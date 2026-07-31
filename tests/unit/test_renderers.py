"""Tests for deterministic Markdown and in-memory DOCX renderers."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from ai_resume_optimizer.models import (
    MatchAnalysis,
    OptimizedResume,
    RequirementAssessment,
    ResumeItem,
    ResumeSection,
)
from ai_resume_optimizer.renderers import (
    render_analysis_report_markdown,
    render_optimized_resume_docx,
    render_optimized_resume_markdown,
)


def _item(
    text: str,
    block_id: str,
    *,
    requirement_id: str | None = None,
    review_note: str | None = None,
) -> ResumeItem:
    return ResumeItem(
        text=text,
        source_block_ids=[block_id],
        related_requirement_ids=[requirement_id] if requirement_id else [],
        needs_review=review_note is not None,
        review_note=review_note,
    )


def _analysis(*, empty_lists: bool = False) -> MatchAnalysis:
    values = [] if empty_lists else ["First item", "Second item"]
    return MatchAnalysis(
        overall_rating="一般",
        overall_evaluation="Relevant evidence exists, with important gaps.",
        assessments=[
            RequirementAssessment(
                requirement_id="requirement-0001",
                status="well_supported",
                source_block_ids=["block-0001"],
                reason="Direct evidence exists.",
                suggested_action="Keep it prominent.",
            ),
            RequirementAssessment(
                requirement_id="requirement-0002",
                status="underrepresented",
                source_block_ids=["block-0002"],
                reason="Related evidence needs context.",
                suggested_action="Clarify the relationship.",
            ),
            RequirementAssessment(
                requirement_id="requirement-0003",
                status="unsupported",
                source_block_ids=[],
                reason="No evidence exists.",
                suggested_action="Do not add it without evidence.",
            ),
        ],
        main_issues=values,
        section_suggestions=values,
        keyword_suggestions=values,
        truthfulness_risks=values,
        content_not_to_add=values,
    )


def _resume(*, include_review_data: bool = True) -> OptimizedResume:
    review_note = "Confirm the scope wording." if include_review_data else None
    return OptimizedResume(
        sections=[
            ResumeSection(
                section_type="basic_info",
                title="Contact Details",
                items=[_item("Jordan Example\njordan@example.invalid", "block-0001")],
                source_block_ids=["block-0001"],
            ),
            ResumeSection(
                section_type="summary",
                title="Profile",
                items=[_item("Backend engineer.", "block-0002")],
                source_block_ids=["block-0002"],
            ),
            ResumeSection(
                section_type="experience",
                title="Selected Experience",
                items=[
                    _item(
                        "Built Python APIs.\nImproved reliability.",
                        "block-0003",
                        requirement_id="requirement-0001",
                        review_note=review_note,
                    ),
                    _item("Maintained services.", "block-0004"),
                ],
                source_block_ids=["block-0003", "block-0004"],
            ),
        ],
        pending_user_inputs=(["Confirm the team size."] if include_review_data else []),
        warnings=["Review all wording."] if include_review_data else [],
    )


def test_analysis_markdown_has_stable_complete_structure() -> None:
    analysis = _analysis()
    before = analysis.model_dump()

    first = render_analysis_report_markdown(analysis)
    second = render_analysis_report_markdown(analysis)

    assert first == second
    assert analysis.model_dump() == before
    assert first.startswith("# AI Resume Optimizer Analysis Report\n")
    assert "**Rating:** 一般" in first
    assert analysis.overall_evaluation in first
    assert first.index("requirement-0001") < first.index("requirement-0002")
    assert first.index("requirement-0002") < first.index("requirement-0003")
    assert "Well Supported" in first
    assert "Underrepresented" in first
    assert "Unsupported" in first
    assert "- **Source Block IDs:** None" in first
    for heading in [
        "## Main Issues",
        "## Section Suggestions",
        "## Keyword Suggestions",
        "## Truthfulness Risks",
        "## Content Not to Add",
        "## Important Notice",
    ]:
        assert heading in first
    assert "not an ATS score or hiring prediction" in first
    assert "%" not in first
    assert first.endswith("\n")
    assert not first.endswith("\n\n")


def test_analysis_markdown_renders_empty_lists_stably() -> None:
    rendered = render_analysis_report_markdown(_analysis(empty_lists=True))

    assert rendered.count("None identified.") == 5


def test_analysis_markdown_preserves_list_order() -> None:
    rendered = render_analysis_report_markdown(_analysis())

    for heading in [
        "## Main Issues",
        "## Section Suggestions",
        "## Keyword Suggestions",
        "## Truthfulness Risks",
        "## Content Not to Add",
    ]:
        section = rendered.split(heading, maxsplit=1)[1]
        assert section.index("First item") < section.index("Second item")


def test_resume_markdown_preserves_sections_items_and_internal_boundaries() -> None:
    resume = _resume()
    before = resume.model_dump()

    first = render_optimized_resume_markdown(resume)
    second = render_optimized_resume_markdown(resume)

    assert first == second
    assert resume.model_dump() == before
    assert first.index("## Contact Details") < first.index("## Profile")
    assert first.index("## Profile") < first.index("## Selected Experience")
    assert "Jordan Example\njordan@example.invalid" in first
    assert "\nBackend engineer.\n" in first
    assert "- Built Python APIs.\n  Improved reliability." in first
    assert first.index("Built Python APIs.") < first.index("Maintained services.")
    assert "block-0001" not in first
    assert "requirement-0001" not in first
    assert "needs_review" not in first
    assert "## Review Notes" in first
    assert "Selected Experience, item 1: Confirm the scope wording." in first
    assert "## Pending User Inputs" in first
    assert "Confirm the team size." in first
    assert "## Warnings" in first
    assert "Review all wording." in first
    assert first.endswith("\n")
    assert not first.endswith("\n\n")


def test_resume_markdown_omits_empty_optional_sections() -> None:
    rendered = render_optimized_resume_markdown(_resume(include_review_data=False))

    assert "## Review Notes" not in rendered
    assert "## Pending User Inputs" not in rendered
    assert "## Warnings" not in rendered


def test_docx_is_valid_ordered_and_uses_expected_builtin_styles() -> None:
    resume = _resume()
    before = resume.model_dump()

    rendered = render_optimized_resume_docx(resume)
    document = Document(BytesIO(rendered))
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs]

    assert isinstance(rendered, bytes)
    assert rendered
    assert resume.model_dump() == before
    assert paragraph_texts[0] == "Optimized Resume"
    assert paragraph_texts.index("Contact Details") < paragraph_texts.index("Profile")
    assert paragraph_texts.index("Profile") < paragraph_texts.index("Selected Experience")
    contact = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Jordan Example")
    )
    summary = next(
        paragraph for paragraph in document.paragraphs if paragraph.text == "Backend engineer."
    )
    experience = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Built Python APIs.")
    )
    maintained = next(
        paragraph for paragraph in document.paragraphs if paragraph.text == "Maintained services."
    )
    assert contact.style.name == "Normal"
    assert summary.style.name == "Normal"
    assert experience.style.name == "List Bullet"
    assert maintained.style.name == "List Bullet"
    assert experience.text == "Built Python APIs.\nImproved reliability."
    assert paragraph_texts.index(experience.text) < paragraph_texts.index(maintained.text)
    full_text = "\n".join(paragraph_texts)
    assert "block-0001" not in full_text
    assert "requirement-0001" not in full_text
    assert "Review Before Use" in paragraph_texts
    assert "Review Notes" in paragraph_texts
    assert "Pending User Inputs" in paragraph_texts
    assert "Warnings" in paragraph_texts


def test_docx_omits_review_appendix_when_no_review_data() -> None:
    rendered = render_optimized_resume_docx(_resume(include_review_data=False))
    document = Document(BytesIO(rendered))
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs]

    assert "Review Before Use" not in paragraph_texts
    assert "Review Notes" not in paragraph_texts
    assert "Pending User Inputs" not in paragraph_texts
    assert "Warnings" not in paragraph_texts
