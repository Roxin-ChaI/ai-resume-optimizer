"""Offline integration-style tests for the single-run pipeline."""

from __future__ import annotations

import hashlib
from collections.abc import Callable
from io import BytesIO
from pathlib import Path
from typing import Never

import pytest
from docx import Document
from docx.document import Document as DocumentObject

import ai_resume_optimizer.pipeline as pipeline_module
import ai_resume_optimizer.runner as runner_module
from ai_resume_optimizer import ResumeOptimizerClosedError, ResumeOptimizerRunner
from ai_resume_optimizer.exceptions import (
    InputError,
    ModelCallError,
    ModelOutputError,
    OutputError,
    ResumeExtractionError,
    TruthfulnessError,
)
from ai_resume_optimizer.models import (
    JobProfile,
    JobRequirement,
    MatchAnalysis,
    OptimizationResult,
    OptimizedResume,
    RequirementAssessment,
    ResumeItem,
    ResumeSection,
    StructuredResume,
)
from ai_resume_optimizer.pipeline import run_optimization
from tests.fakes import FakeModelClient

DocxFactory = Callable[[Callable[[DocumentObject], None] | None, str], Path]
PdfFactory = Callable[[list[str], str], Path]


class _ClosableFakeModelClient(FakeModelClient):
    def __init__(self) -> None:
        super().__init__()
        self.close_calls = 0

    def close(self) -> None:
        self.close_calls += 1


def _resume_document(document: DocumentObject) -> None:
    document.add_paragraph("Built Python APIs in 2023.")


def _structured_resume() -> StructuredResume:
    item = ResumeItem(
        text="Built Python APIs in 2023.",
        source_block_ids=["block-0001"],
        related_requirement_ids=[],
        needs_review=False,
        review_note=None,
    )
    return StructuredResume(
        sections=[
            ResumeSection(
                section_type="experience",
                title="Experience",
                items=[item],
                source_block_ids=["block-0001"],
            )
        ],
        unclassified_content=[],
        warnings=[],
    )


def _job_profile() -> JobProfile:
    return JobProfile(
        role_summary="Build backend services.",
        requirements=[
            JobRequirement(
                requirement_id="requirement-0001",
                category="core_skill",
                description="Python",
                importance="required",
                source_excerpt="Python",
            )
        ],
    )


def _match_analysis() -> MatchAnalysis:
    return MatchAnalysis(
        overall_rating="高",
        overall_evaluation="The resume contains direct supporting evidence.",
        assessments=[
            RequirementAssessment(
                requirement_id="requirement-0001",
                status="well_supported",
                source_block_ids=["block-0001"],
                reason="Python API experience is explicit.",
                suggested_action="Keep the evidence prominent.",
            )
        ],
        main_issues=[],
        section_suggestions=[],
        keyword_suggestions=[],
        truthfulness_risks=[],
        content_not_to_add=[],
    )


def _optimized_resume() -> OptimizedResume:
    item = ResumeItem(
        text="Built Python APIs in 2023.",
        source_block_ids=["block-0001"],
        related_requirement_ids=["requirement-0001"],
        needs_review=False,
        review_note=None,
    )
    return OptimizedResume(
        sections=[
            ResumeSection(
                section_type="experience",
                title="Experience",
                items=[item],
                source_block_ids=["block-0001"],
            )
        ],
        pending_user_inputs=[],
        warnings=["duplicate warning", "optimizer warning"],
    )


def _fake_client() -> FakeModelClient:
    return FakeModelClient(
        {
            StructuredResume: _structured_resume(),
            JobProfile: _job_profile(),
            MatchAnalysis: _match_analysis(),
            OptimizedResume: _optimized_resume(),
        }
    )


def _file_digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _final_paths(output_dir: Path) -> list[Path]:
    return [
        output_dir / "analysis_report.md",
        output_dir / "optimized_resume.md",
        output_dir / "optimized_resume.docx",
    ]


def _assert_no_output_batch(output_dir: Path) -> None:
    assert not any(path.exists() for path in _final_paths(output_dir))
    if output_dir.exists() and output_dir.is_dir():
        assert not list(output_dir.glob("*.tmp"))


def test_in_memory_core_returns_structured_result_without_rendering_or_writing(
    tmp_path: Path,
    docx_factory: DocxFactory,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    resume_path = docx_factory(_resume_document)
    client = _fake_client()

    def fail(*args: object, **kwargs: object) -> Never:
        raise AssertionError("In-memory optimization must not render or write outputs.")

    monkeypatch.setattr(pipeline_module, "_render_outputs", fail)
    monkeypatch.setattr(pipeline_module, "_write_output_batch", fail)

    result = pipeline_module._run_optimization_in_memory(
        resume_path=resume_path,
        job_description="  Python is required.  ",
        model_client=client,
    )

    assert isinstance(result, OptimizationResult)
    assert result.output_paths == {}
    assert result.analysis == _match_analysis()
    assert result.optimized_resume == _optimized_resume()
    assert result.warnings == ["duplicate warning", "optimizer warning"]
    assert [call.response_model for call in client.calls] == [
        StructuredResume,
        JobProfile,
        MatchAnalysis,
        OptimizedResume,
    ]
    assert client.calls[1].input_text == "Python is required."
    assert list(tmp_path.iterdir()) == [resume_path]


def test_runner_optimizes_docx_in_memory(
    tmp_path: Path,
    docx_factory: DocxFactory,
) -> None:
    resume_path = docx_factory(_resume_document)
    client = _fake_client()

    result = ResumeOptimizerRunner(client).optimize(
        resume_path=resume_path,
        job_description="Python is required.",
    )

    assert isinstance(result, OptimizationResult)
    assert result.output_paths == {}
    assert result.analysis == _match_analysis()
    assert result.optimized_resume == _optimized_resume()
    assert list(tmp_path.iterdir()) == [resume_path]


def test_runner_optimizes_pdf_in_memory(
    tmp_path: Path,
    pdf_factory: PdfFactory,
) -> None:
    resume_path = pdf_factory(["Built Python APIs in 2023."])
    client = _fake_client()

    result = ResumeOptimizerRunner(client).optimize(
        resume_path=resume_path,
        job_description="Python is required.",
    )

    assert isinstance(result, OptimizationResult)
    assert result.output_paths == {}
    assert result.analysis == _match_analysis()
    assert result.optimized_resume == _optimized_resume()
    assert list(tmp_path.iterdir()) == [resume_path]


def test_runner_forwards_exact_inputs_to_shared_core(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    resume_path = tmp_path / "resume.docx"
    job_description = "  Python is required.  "
    client = FakeModelClient()
    expected = OptimizationResult(
        analysis=_match_analysis(),
        optimized_resume=_optimized_resume(),
        output_paths={},
        warnings=[],
    )
    received: dict[str, object] = {}

    def run_in_memory(**kwargs: object) -> OptimizationResult:
        received.update(kwargs)
        return expected

    monkeypatch.setattr(runner_module, "_run_optimization_in_memory", run_in_memory)

    actual = ResumeOptimizerRunner(client).optimize(
        resume_path=resume_path,
        job_description=job_description,
    )

    assert actual is expected
    assert received == {
        "resume_path": resume_path,
        "job_description": job_description,
        "model_client": client,
    }
    assert not resume_path.exists()


def test_runner_close_is_idempotent_and_does_not_close_external_client() -> None:
    client = _ClosableFakeModelClient()
    runner = ResumeOptimizerRunner(client)

    runner.close()
    runner.close()

    assert client.close_calls == 0


def test_runner_close_closes_owned_client_once() -> None:
    client = _ClosableFakeModelClient()
    runner = ResumeOptimizerRunner(client, owns_model_client=True)

    runner.close()
    runner.close()

    assert client.close_calls == 1


def test_runner_rejects_owned_client_without_close() -> None:
    with pytest.raises(TypeError, match="must implement close"):
        ResumeOptimizerRunner(FakeModelClient(), owns_model_client=True)


def test_closed_runner_rejects_optimization_before_any_work(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    runner = ResumeOptimizerRunner(FakeModelClient())
    runner.close()

    def fail(*args: object, **kwargs: object) -> Never:
        raise AssertionError("Closed runner must not enter the optimization core.")

    monkeypatch.setattr(runner_module, "_run_optimization_in_memory", fail)

    with pytest.raises(ResumeOptimizerClosedError, match="runner is closed"):
        runner.optimize(
            resume_path=tmp_path / "resume.docx",
            job_description="Python is required.",
        )

    assert not list(tmp_path.iterdir())


def test_pipeline_runs_complete_ordered_flow_and_writes_one_consistent_batch(
    tmp_path: Path,
    docx_factory: DocxFactory,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    resume_path = docx_factory(_resume_document)
    original_digest = _file_digest(resume_path)
    output_dir = tmp_path / "nested" / "output"
    client = _fake_client()
    original_parse_resume = pipeline_module.parse_resume

    def parse_with_warnings(path: Path):
        extracted = original_parse_resume(path)
        return extracted.model_copy(update={"warnings": ["parser warning", "duplicate warning"]})

    monkeypatch.setattr(pipeline_module, "parse_resume", parse_with_warnings)

    result = run_optimization(
        resume_path=resume_path,
        job_description="  Python is required.  ",
        output_dir=output_dir,
        model_client=client,
    )

    assert [call.response_model for call in client.calls] == [
        StructuredResume,
        JobProfile,
        MatchAnalysis,
        OptimizedResume,
    ]
    assert client.calls[1].input_text == "Python is required."
    assert set(result.output_paths) == {
        "analysis_report",
        "optimized_resume_markdown",
        "optimized_resume_docx",
    }
    assert result.analysis == _match_analysis()
    assert result.optimized_resume == _optimized_resume()
    assert result.warnings == [
        "parser warning",
        "duplicate warning",
        "optimizer warning",
    ]
    assert [path.name for path in result.output_paths.values()] == [
        "analysis_report.md",
        "optimized_resume.md",
        "optimized_resume.docx",
    ]
    analysis_text = result.output_paths["analysis_report"].read_text(encoding="utf-8")
    resume_text = result.output_paths["optimized_resume_markdown"].read_text(encoding="utf-8")
    docx_bytes = result.output_paths["optimized_resume_docx"].read_bytes()
    document = Document(BytesIO(docx_bytes))
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert analysis_text
    assert resume_text
    assert docx_bytes
    assert "Built Python APIs in 2023." in resume_text
    assert "Built Python APIs in 2023." in document_text
    assert _file_digest(resume_path) == original_digest


def test_pipeline_succeeds_with_existing_empty_output_directory(
    tmp_path: Path,
    docx_factory: DocxFactory,
) -> None:
    output_dir = tmp_path / "output"
    output_dir.mkdir()

    run_optimization(
        resume_path=docx_factory(_resume_document),
        job_description="Python is required.",
        output_dir=output_dir,
        model_client=_fake_client(),
    )

    assert all(path.exists() for path in _final_paths(output_dir))


def test_pipeline_rejects_output_directory_that_is_a_file(
    tmp_path: Path,
) -> None:
    output_path = tmp_path / "output"
    output_path.write_text("user data", encoding="utf-8")
    client = _fake_client()

    with pytest.raises(OutputError):
        run_optimization(
            resume_path=tmp_path / "unused.docx",
            job_description="Python is required.",
            output_dir=output_path,
            model_client=client,
        )

    assert output_path.read_text(encoding="utf-8") == "user data"
    assert client.calls == []


@pytest.mark.parametrize(
    "filename",
    ["analysis_report.md", "optimized_resume.md", "optimized_resume.docx"],
)
def test_pipeline_refuses_overwrite_before_model_calls(
    tmp_path: Path,
    filename: str,
) -> None:
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    existing = output_dir / filename
    existing.write_bytes(b"existing user content")
    client = _fake_client()

    with pytest.raises(OutputError, match="Refusing to overwrite"):
        run_optimization(
            resume_path=tmp_path / "unused.docx",
            job_description="Python is required.",
            output_dir=output_dir,
            model_client=client,
        )

    assert existing.read_bytes() == b"existing user content"
    assert client.calls == []


def test_pipeline_rejects_empty_job_description_before_model_calls(
    tmp_path: Path,
) -> None:
    client = _fake_client()
    output_dir = tmp_path / "output"

    with pytest.raises(InputError):
        run_optimization(
            resume_path=tmp_path / "unused.docx",
            job_description=" \n ",
            output_dir=output_dir,
            model_client=client,
        )

    assert client.calls == []
    _assert_no_output_batch(output_dir)


def test_pipeline_parser_failure_leaves_no_final_files(tmp_path: Path) -> None:
    output_dir = tmp_path / "output"
    corrupt_resume = tmp_path / "corrupt.docx"
    corrupt_resume.write_bytes(b"not a DOCX file")

    with pytest.raises(ResumeExtractionError):
        run_optimization(
            resume_path=corrupt_resume,
            job_description="Python is required.",
            output_dir=output_dir,
            model_client=_fake_client(),
        )

    _assert_no_output_batch(output_dir)


def test_pipeline_model_call_failure_keeps_original_error(
    tmp_path: Path,
    docx_factory: DocxFactory,
) -> None:
    output_dir = tmp_path / "output"
    expected = ModelCallError("provider unavailable")

    with pytest.raises(ModelCallError) as raised:
        run_optimization(
            resume_path=docx_factory(_resume_document),
            job_description="Python is required.",
            output_dir=output_dir,
            model_client=FakeModelClient(error=expected),
        )

    assert raised.value is expected
    _assert_no_output_batch(output_dir)


@pytest.mark.parametrize(
    ("stage_name", "error"),
    [
        ("analyze_job", ModelOutputError("invalid job output")),
        ("analyze_match", ModelOutputError("invalid match output")),
        ("optimize_resume", ModelOutputError("invalid optimization output")),
        ("validate_optimized_resume", TruthfulnessError("unsupported fact")),
    ],
)
def test_pipeline_business_stage_failures_keep_original_error_and_no_outputs(
    tmp_path: Path,
    docx_factory: DocxFactory,
    monkeypatch: pytest.MonkeyPatch,
    stage_name: str,
    error: Exception,
) -> None:
    output_dir = tmp_path / "output"

    def fail(*args: object, **kwargs: object) -> Never:
        raise error

    monkeypatch.setattr(pipeline_module, stage_name, fail)

    with pytest.raises(type(error)) as raised:
        run_optimization(
            resume_path=docx_factory(_resume_document),
            job_description="Python is required.",
            output_dir=output_dir,
            model_client=_fake_client(),
        )

    assert raised.value is error
    _assert_no_output_batch(output_dir)


@pytest.mark.parametrize(
    "renderer_name",
    [
        "render_analysis_report_markdown",
        "render_optimized_resume_markdown",
        "render_optimized_resume_docx",
    ],
)
def test_pipeline_renderer_failures_become_output_errors_without_files(
    tmp_path: Path,
    docx_factory: DocxFactory,
    monkeypatch: pytest.MonkeyPatch,
    renderer_name: str,
) -> None:
    output_dir = tmp_path / "output"

    def fail(*args: object, **kwargs: object) -> Never:
        raise OSError("synthetic rendering failure")

    monkeypatch.setattr(pipeline_module, renderer_name, fail)

    with pytest.raises(OutputError) as raised:
        run_optimization(
            resume_path=docx_factory(_resume_document),
            job_description="Python is required.",
            output_dir=output_dir,
            model_client=_fake_client(),
        )

    assert isinstance(raised.value.__cause__, OSError)
    _assert_no_output_batch(output_dir)


@pytest.mark.parametrize(
    "failed_name",
    ["optimized_resume.md", "optimized_resume.docx"],
)
def test_pipeline_temporary_write_failure_cleans_batch(
    tmp_path: Path,
    docx_factory: DocxFactory,
    monkeypatch: pytest.MonkeyPatch,
    failed_name: str,
) -> None:
    output_dir = tmp_path / "output"
    original_writer = pipeline_module._write_temporary_file

    def fail_selected_write(
        directory: Path,
        final_name: str,
        content: bytes,
    ) -> Path:
        if final_name == failed_name:
            raise OSError("synthetic temporary write failure")
        return original_writer(directory, final_name, content)

    monkeypatch.setattr(
        pipeline_module,
        "_write_temporary_file",
        fail_selected_write,
    )

    with pytest.raises(OutputError):
        run_optimization(
            resume_path=docx_factory(_resume_document),
            job_description="Python is required.",
            output_dir=output_dir,
            model_client=_fake_client(),
        )

    _assert_no_output_batch(output_dir)


def test_pipeline_mid_move_failure_removes_created_batch_files(
    tmp_path: Path,
    docx_factory: DocxFactory,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    output_dir = tmp_path / "output"
    original_replace = pipeline_module._replace_file
    call_count = 0

    def fail_second_move(source: Path, target: Path) -> None:
        nonlocal call_count
        call_count += 1
        if call_count == 2:
            raise OSError("synthetic move failure")
        original_replace(source, target)

    monkeypatch.setattr(pipeline_module, "_replace_file", fail_second_move)

    with pytest.raises(OutputError):
        run_optimization(
            resume_path=docx_factory(_resume_document),
            job_description="Python is required.",
            output_dir=output_dir,
            model_client=_fake_client(),
        )

    _assert_no_output_batch(output_dir)
