"""Tests for PDF, DOCX, and job-description input parsing."""

from __future__ import annotations

import codecs
import hashlib
from collections.abc import Callable
from pathlib import Path

import pytest
from docx.document import Document as DocumentObject

import ai_resume_optimizer.parsers.docx as docx_parser_module
import ai_resume_optimizer.parsers.job_description as job_parser_module
import ai_resume_optimizer.parsers.pdf as pdf_parser_module
from ai_resume_optimizer.exceptions import (
    InputError,
    InputTooLargeError,
    ResumeExtractionError,
    UnsupportedFormatError,
)
from ai_resume_optimizer.parsers import parse_resume
from ai_resume_optimizer.parsers.docx import parse_docx_resume
from ai_resume_optimizer.parsers.job_description import (
    normalize_job_description,
    read_job_description,
)
from ai_resume_optimizer.parsers.pdf import parse_pdf_resume

PdfFactory = Callable[[list[str], str], Path]
DocxFactory = Callable[[Callable[[DocumentObject], None] | None, str], Path]
TextFileFactory = Callable[[str | bytes, str], Path]


def _file_state(path: Path) -> tuple[str, int]:
    return hashlib.sha256(path.read_bytes()).hexdigest(), path.stat().st_mtime_ns


def _simple_docx(document: DocumentObject) -> None:
    document.add_paragraph("Jordan Example has extensive Python engineering experience.")


def test_parse_resume_dispatches_supported_formats_case_insensitively(
    pdf_factory: PdfFactory,
    docx_factory: DocxFactory,
) -> None:
    pdf_path = pdf_factory(
        ["Jordan Example has extensive Python engineering experience."],
        "RESUME.PDF",
    )
    docx_path = docx_factory(_simple_docx, "RESUME.DOCX")

    assert parse_resume(pdf_path).source_format == "pdf"
    assert parse_resume(docx_path).source_format == "docx"


def test_parse_resume_rejects_unsupported_existing_file(
    text_file_factory: TextFileFactory,
) -> None:
    path = text_file_factory("resume content", "resume.rtf")

    with pytest.raises(UnsupportedFormatError, match=str(path)):
        parse_resume(path)


def test_parse_resume_rejects_missing_file_and_directory(tmp_path: Path) -> None:
    with pytest.raises(InputError, match="does not exist"):
        parse_resume(tmp_path / "missing.pdf")

    with pytest.raises(InputError, match="not a regular file"):
        parse_resume(tmp_path)


def test_pdf_extracts_multiple_pages_in_stable_order(
    pdf_factory: PdfFactory,
) -> None:
    path = pdf_factory(
        [
            "Jordan Example builds reliable Python backend services.",
            "Second page describes cloud deployment and monitoring experience.",
        ]
    )
    before = _file_state(path)

    first = parse_pdf_resume(path)
    second = parse_pdf_resume(path)

    assert first.source_format == "pdf"
    assert first.source_path == path
    assert [block.text for block in first.blocks] == [
        "Jordan Example builds reliable Python backend services.",
        "Second page describes cloud deployment and monitoring experience.",
    ]
    assert [block.block_id for block in first.blocks] == ["block-0001", "block-0002"]
    assert [block.location for block in first.blocks] == ["page:1", "page:2"]
    assert all(block.kind == "paragraph" for block in first.blocks)
    assert first.plain_text == "\n\n".join(block.text for block in first.blocks)
    assert first == second
    assert _file_state(path) == before


@pytest.mark.parametrize("pages", [[], [""]])
def test_pdf_without_extractable_text_reports_ocr_limit(
    pdf_factory: PdfFactory,
    pages: list[str],
) -> None:
    path = pdf_factory(pages)

    with pytest.raises(ResumeExtractionError, match="OCR") as error:
        parse_pdf_resume(path)

    assert "Scanned PDFs" in str(error.value)


def test_corrupt_pdf_preserves_exception_chain(tmp_path: Path) -> None:
    path = tmp_path / "corrupt.pdf"
    path.write_bytes(b"not a pdf")

    with pytest.raises(ResumeExtractionError) as error:
        parse_pdf_resume(path)

    assert error.value.__cause__ is not None


def test_pdf_enforces_normalized_character_limit(
    pdf_factory: PdfFactory,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = pdf_factory(["Jordan Example has more than thirty normalized characters."])
    monkeypatch.setattr(pdf_parser_module, "MAX_RESUME_CHARACTERS", 30)

    with pytest.raises(InputTooLargeError, match="30") as error:
        parse_pdf_resume(path)

    assert "will not be truncated" in str(error.value)


def test_direct_pdf_parser_rejects_wrong_suffix(
    text_file_factory: TextFileFactory,
) -> None:
    path = text_file_factory("not a PDF", "resume.txt")

    with pytest.raises(UnsupportedFormatError, match=str(path)):
        parse_pdf_resume(path)


def _structured_docx(document: DocumentObject) -> None:
    document.add_heading("Experience", level=1)
    document.add_paragraph("Built reliable services")
    document.add_paragraph("Python", style="List Bullet")

    numbered = document.add_paragraph("Explicit numbering")
    numbered_properties = numbered._p.get_or_add_pPr().get_or_add_numPr()
    numbered_properties.get_or_add_numId().val = 1

    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Company"
    table.rows[0].cells[0].add_paragraph("Platform Team")
    table.rows[0].cells[1].text = "Software Engineer"
    table.rows[0].cells[2].text = "2022-2025"
    table.rows[1].cells[0].text = "Project"
    table.rows[1].cells[1].text = "API modernization"
    table.rows[1].cells[2].text = "Python"

    document.add_paragraph("Education details")


def test_docx_extracts_structures_and_preserves_body_order(
    docx_factory: DocxFactory,
) -> None:
    path = docx_factory(_structured_docx)
    before = _file_state(path)

    result = parse_docx_resume(path)

    assert result.source_format == "docx"
    assert result.source_path == path
    assert [block.kind for block in result.blocks] == [
        "heading",
        "paragraph",
        "list_item",
        "list_item",
        "table_row",
        "table_row",
        "paragraph",
    ]
    assert [block.text for block in result.blocks] == [
        "Experience",
        "Built reliable services",
        "Python",
        "Explicit numbering",
        "Company / Platform Team | Software Engineer | 2022-2025",
        "Project | API modernization | Python",
        "Education details",
    ]
    assert [block.block_id for block in result.blocks] == [
        f"block-{number:04d}" for number in range(1, 8)
    ]
    assert result.blocks[4].location == "table:1,row:1"
    assert result.blocks[5].location == "table:1,row:2"
    assert result.plain_text == "\n\n".join(block.text for block in result.blocks)
    assert _file_state(path) == before


def _merged_cells_docx(document: DocumentObject) -> None:
    table = document.add_table(rows=2, cols=2)
    merged = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    merged.text = "Merged value"
    table.rows[1].cells[0].text = "Same value"
    table.rows[1].cells[1].text = "Same value"


def test_docx_deduplicates_only_the_same_merged_cell(
    docx_factory: DocxFactory,
) -> None:
    path = docx_factory(_merged_cells_docx)

    result = parse_docx_resume(path)

    assert [block.text for block in result.blocks] == [
        "Merged value",
        "Same value | Same value",
    ]


def test_empty_docx_is_rejected(docx_factory: DocxFactory) -> None:
    path = docx_factory()

    with pytest.raises(ResumeExtractionError, match="does not contain"):
        parse_docx_resume(path)


def test_corrupt_docx_preserves_exception_chain(tmp_path: Path) -> None:
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"not a docx")

    with pytest.raises(ResumeExtractionError) as error:
        parse_docx_resume(path)

    assert error.value.__cause__ is not None


def test_docx_enforces_normalized_character_limit(
    docx_factory: DocxFactory,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = docx_factory(_simple_docx)
    monkeypatch.setattr(docx_parser_module, "MAX_RESUME_CHARACTERS", 20)

    with pytest.raises(InputTooLargeError, match="20") as error:
        parse_docx_resume(path)

    assert "will not be truncated" in str(error.value)


def test_direct_docx_parser_rejects_wrong_suffix(
    text_file_factory: TextFileFactory,
) -> None:
    path = text_file_factory("not a DOCX", "resume.txt")

    with pytest.raises(UnsupportedFormatError, match=str(path)):
        parse_docx_resume(path)


def test_normalize_job_description_is_deterministic() -> None:
    text = "  Senior   Engineer  \r\n\r\n\r\nPython\x00  \rNext\t\tline  "

    assert normalize_job_description(text) == "Senior Engineer\n\nPython\nNext line"


def test_job_description_file_and_direct_text_share_normalization(
    text_file_factory: TextFileFactory,
) -> None:
    raw_text = "Role  \r\n\r\n\r\nPython   experience  "
    path = text_file_factory(raw_text)
    before = _file_state(path)

    assert read_job_description(path) == normalize_job_description(raw_text)
    assert _file_state(path) == before


def test_job_description_reads_utf8_bom_case_insensitively(
    text_file_factory: TextFileFactory,
) -> None:
    path = text_file_factory(
        codecs.BOM_UTF8 + "Python engineer".encode(),
        "JOB.TXT",
    )

    assert read_job_description(path) == "Python engineer"


@pytest.mark.parametrize("text", ["", " ", "\x00\r\n\t"])
def test_empty_job_description_text_is_rejected(text: str) -> None:
    with pytest.raises(InputError, match="empty"):
        normalize_job_description(text)


def test_empty_job_description_file_is_rejected(
    text_file_factory: TextFileFactory,
) -> None:
    path = text_file_factory("")

    with pytest.raises(InputError, match="empty"):
        read_job_description(path)


def test_non_utf8_job_description_preserves_exception_chain(
    text_file_factory: TextFileFactory,
) -> None:
    path = text_file_factory(b"\xff\xfe\xfa")

    with pytest.raises(InputError) as error:
        read_job_description(path)

    assert error.value.__cause__ is not None


def test_job_description_rejects_wrong_suffix(
    text_file_factory: TextFileFactory,
) -> None:
    path = text_file_factory("Python engineer", "job.md")

    with pytest.raises(UnsupportedFormatError, match=str(path)):
        read_job_description(path)


def test_job_description_rejects_missing_file_and_directory(tmp_path: Path) -> None:
    with pytest.raises(InputError, match="does not exist"):
        read_job_description(tmp_path / "missing.txt")

    directory = tmp_path / "directory.txt"
    directory.mkdir()
    with pytest.raises(InputError, match="not a regular file"):
        read_job_description(directory)


def test_job_description_enforces_normalized_character_limit(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(job_parser_module, "MAX_JOB_DESCRIPTION_CHARACTERS", 10)

    with pytest.raises(InputTooLargeError, match="10") as error:
        normalize_job_description("Senior Python engineer")

    assert "will not be truncated" in str(error.value)
