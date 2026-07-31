"""In-memory DOCX renderer for optimized resumes."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph

from ai_resume_optimizer.models import OptimizedResume

_BULLETED_SECTION_TYPES = {
    "skills",
    "experience",
    "projects",
    "education",
    "certifications",
    "other",
}


def _set_multiline_text(paragraph: Paragraph, text: str) -> None:
    lines = text.splitlines()
    paragraph.add_run(lines[0])
    for line in lines[1:]:
        paragraph.add_run().add_break()
        paragraph.add_run(line)


def _add_bullet(document: DocumentObject, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    _set_multiline_text(paragraph, text)


def render_optimized_resume(resume: OptimizedResume) -> bytes:
    """Render an editable DOCX entirely in memory."""

    document = Document()
    document.add_heading("Optimized Resume", level=0)
    review_notes: list[str] = []

    for section in resume.sections:
        document.add_heading(section.title, level=1)
        is_bulleted = section.section_type in _BULLETED_SECTION_TYPES
        for item_number, item in enumerate(section.items, start=1):
            if is_bulleted:
                paragraph = document.add_paragraph(style="List Bullet")
            else:
                paragraph = document.add_paragraph()
            _set_multiline_text(paragraph, item.text)
            if item.needs_review:
                review_notes.append(f"{section.title}, item {item_number}: {item.review_note}")

    if review_notes or resume.pending_user_inputs or resume.warnings:
        document.add_page_break()
        document.add_heading("Review Before Use", level=1)
        if review_notes:
            document.add_heading("Review Notes", level=2)
            for note in review_notes:
                _add_bullet(document, note)
        if resume.pending_user_inputs:
            document.add_heading("Pending User Inputs", level=2)
            for pending_input in resume.pending_user_inputs:
                _add_bullet(document, pending_input)
        if resume.warnings:
            document.add_heading("Warnings", level=2)
            for warning in resume.warnings:
                _add_bullet(document, warning)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
